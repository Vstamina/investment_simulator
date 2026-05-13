import streamlit as st
import pandas as pd
import numpy as np
import requests
import yfinance as yf
import plotly.graph_objects as go

from io import BytesIO
from docx import Document
from datetime import date
from docx.shared import Inches


# ============================================================
# 1. COLETA DE DADOS
# ============================================================

@st.cache_data(ttl=60 * 60 * 12)
def carregar_ibovespa(data_inicio: str, data_fim: str) -> pd.DataFrame:
    """
    Carrega o histórico do Ibovespa pelo Yahoo Finance.
    Ticker: ^BVSP.

    Esta versão trata o caso em que o yfinance retorna colunas em MultiIndex.
    """

    df = yf.download(
        "^BVSP",
        start=data_inicio,
        end=data_fim,
        progress=False,
        auto_adjust=True
    )

    if df.empty:
        return pd.DataFrame()

    # Corrige quando o yfinance retorna colunas com múltiplos níveis
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [
            col[0] if isinstance(col, tuple) else col
            for col in df.columns
        ]

    df = df.reset_index()

    # Padroniza a coluna de data
    if "Date" in df.columns:
        df = df.rename(columns={"Date": "data"})
    elif "Datetime" in df.columns:
        df = df.rename(columns={"Datetime": "data"})

    # Padroniza a coluna de preço
    if "Close" in df.columns:
        df = df.rename(columns={"Close": "ibovespa"})
    elif "Adj Close" in df.columns:
        df = df.rename(columns={"Adj Close": "ibovespa"})
    elif "close" in df.columns:
        df = df.rename(columns={"close": "ibovespa"})

    if "data" not in df.columns or "ibovespa" not in df.columns:
        return pd.DataFrame()

    df = df[["data", "ibovespa"]].copy()
    df["data"] = pd.to_datetime(df["data"], errors="coerce")
    df["ibovespa"] = pd.to_numeric(df["ibovespa"], errors="coerce")

    df = df.dropna()
    df = df.drop_duplicates(subset=["data"])
    df = df.sort_values("data")

    return df


@st.cache_data(ttl=60 * 60 * 12)
def carregar_cdi_bacen(data_inicio: str, data_fim: str) -> pd.DataFrame:
    """
    Carrega CDI diário pelo SGS/Bacen.
    Série principal: 12.
    Fallback: série 11, caso a série 12 falhe no ambiente.
    Faz leitura em blocos anuais para evitar falhas em períodos longos.
    """

    inicio = pd.to_datetime(data_inicio)
    fim = pd.to_datetime(data_fim)

    if inicio >= fim:
        return pd.DataFrame()

    series_tentativa = [12, 11]
    todos_dados = []

    for codigo_serie in series_tentativa:
        todos_dados = []

        ano_inicio = inicio.year
        ano_fim = fim.year

        for ano in range(ano_inicio, ano_fim + 1):
            bloco_inicio = max(inicio, pd.Timestamp(year=ano, month=1, day=1))
            bloco_fim = min(fim, pd.Timestamp(year=ano, month=12, day=31))

            data_inicio_br = bloco_inicio.strftime("%d/%m/%Y")
            data_fim_br = bloco_fim.strftime("%d/%m/%Y")

            url = (
                f"https://api.bcb.gov.br/dados/serie/bcdata.sgs.{codigo_serie}/dados"
                f"?formato=json&dataInicial={data_inicio_br}&dataFinal={data_fim_br}"
            )

            try:
                response = requests.get(url, timeout=30)
                response.raise_for_status()
                dados = response.json()

                if dados:
                    todos_dados.extend(dados)

            except Exception:
                continue

        if todos_dados:
            df = pd.DataFrame(todos_dados)

            if df.empty:
                continue

            df["data"] = pd.to_datetime(df["data"], dayfirst=True, errors="coerce")
            df["cdi_diario_pct"] = (
                df["valor"]
                .astype(str)
                .str.replace(",", ".", regex=False)
                .astype(float)
            )

            df = df[["data", "cdi_diario_pct"]].copy()
            df = df.dropna()
            df = df.drop_duplicates(subset=["data"])
            df = df.sort_values("data")

            return df

    return pd.DataFrame()


# ============================================================
# 2. PREPARAÇÃO E MÉTRICAS
# ============================================================

def preparar_base_comparativa(df_ibov: pd.DataFrame, df_cdi: pd.DataFrame) -> pd.DataFrame:
    """
    Une Ibovespa e CDI por data e calcula índices acumulados.
    Base 100 para comparação visual.
    """

    if df_ibov.empty or df_cdi.empty:
        return pd.DataFrame()

    df = pd.merge(df_ibov, df_cdi, on="data", how="inner")
    df = df.sort_values("data").copy()

    if df.empty:
        return pd.DataFrame()

    df["retorno_ibov_diario"] = df["ibovespa"].pct_change().fillna(0)
    df["fator_cdi_diario"] = 1 + (df["cdi_diario_pct"] / 100)

    df["indice_ibov"] = 100 * (1 + df["retorno_ibov_diario"]).cumprod()
    df["indice_cdi"] = 100 * df["fator_cdi_diario"].cumprod() / df["fator_cdi_diario"].iloc[0]

    df["retorno_acum_ibov"] = (df["indice_ibov"] / 100) - 1
    df["retorno_acum_cdi"] = (df["indice_cdi"] / 100) - 1

    return df


def calcular_metricas(df: pd.DataFrame) -> dict:
    """
    Calcula métricas principais da comparação.
    """

    if df.empty:
        return {}

    data_inicio = df["data"].min()
    data_fim = df["data"].max()

    dias_corridos = max((data_fim - data_inicio).days, 1)
    anos = dias_corridos / 365.25
    meses = dias_corridos / 30.4375

    retorno_ibov = df["indice_ibov"].iloc[-1] / df["indice_ibov"].iloc[0] - 1
    retorno_cdi = df["indice_cdi"].iloc[-1] / df["indice_cdi"].iloc[0] - 1

    retorno_anual_ibov = (1 + retorno_ibov) ** (1 / anos) - 1
    retorno_anual_cdi = (1 + retorno_cdi) ** (1 / anos) - 1

    retorno_mensal_ibov = (1 + retorno_ibov) ** (1 / meses) - 1
    retorno_mensal_cdi = (1 + retorno_cdi) ** (1 / meses) - 1

    volatilidade_ibov = df["retorno_ibov_diario"].std() * np.sqrt(252)

    df_temp = df.copy()
    df_temp["max_ibov"] = df_temp["indice_ibov"].cummax()
    df_temp["drawdown_ibov"] = df_temp["indice_ibov"] / df_temp["max_ibov"] - 1
    max_drawdown_ibov = df_temp["drawdown_ibov"].min()

    vencedor = "Ibovespa" if retorno_ibov > retorno_cdi else "CDI"

    return {
        "data_inicio": data_inicio,
        "data_fim": data_fim,
        "anos": anos,
        "retorno_ibov": retorno_ibov,
        "retorno_cdi": retorno_cdi,
        "retorno_anual_ibov": retorno_anual_ibov,
        "retorno_anual_cdi": retorno_anual_cdi,
        "retorno_mensal_ibov": retorno_mensal_ibov,
        "retorno_mensal_cdi": retorno_mensal_cdi,
        "volatilidade_ibov": volatilidade_ibov,
        "max_drawdown_ibov": max_drawdown_ibov,
        "vencedor": vencedor,
    }


def formatar_pct(valor: float) -> str:
    return f"{valor * 100:,.2f}%".replace(",", "X").replace(".", ",").replace("X", ".")


def formatar_numero(valor: float) -> str:
    return f"{valor:,.0f}".replace(",", ".")


# ============================================================
# 3. DETECÇÃO SIMPLIFICADA DE CICLOS
# ============================================================

def detectar_ciclos_relevantes(
    df: pd.DataFrame,
    numero_maximo_ciclos: int = 4,
    queda_minima: float = 0.20,
    alta_minima: float = 0.35,
    meses_minimos: int = 12
) -> pd.DataFrame:
    """
    Detecta janelas longas em que houve queda relevante seguida de recuperação.
    A lógica é simplificada e consultiva, não é recomendação automática.
    """

    if df.empty:
        return pd.DataFrame()

    base = df.copy()
    base = base.sort_values("data")
    base["max_anterior"] = base["ibovespa"].cummax()
    base["queda_desde_topo"] = base["ibovespa"] / base["max_anterior"] - 1

    fundos = base[base["queda_desde_topo"] <= -queda_minima].copy()

    if fundos.empty:
        return pd.DataFrame()

    ciclos = []
    datas_usadas = []

    for _, fundo in fundos.iterrows():
        data_entrada = fundo["data"]
        ibov_entrada = fundo["ibovespa"]

        if any(abs((data_entrada - d).days) < 180 for d in datas_usadas):
            continue

        data_minima_saida = data_entrada + pd.DateOffset(months=meses_minimos)

        posteriores = base[
            (base["data"] >= data_minima_saida)
            & (base["ibovespa"] >= ibov_entrada * (1 + alta_minima))
        ].copy()

        if posteriores.empty:
            continue

        saida = posteriores.iloc[0]
        data_saida = saida["data"]
        ibov_saida = saida["ibovespa"]

        janela = base[
            (base["data"] >= data_entrada)
            & (base["data"] <= data_saida)
        ].copy()

        if janela.empty:
            continue

        retorno_ibov = ibov_saida / ibov_entrada - 1
        retorno_cdi = janela["indice_cdi"].iloc[-1] / janela["indice_cdi"].iloc[0] - 1

        vencedor = "Ibovespa" if retorno_ibov > retorno_cdi else "CDI"

        ciclos.append({
            "entrada": data_entrada.date(),
            "saida": data_saida.date(),
            "ibov_entrada": ibov_entrada,
            "ibov_saida": ibov_saida,
            "retorno_ibov": retorno_ibov,
            "retorno_cdi": retorno_cdi,
            "diferenca": retorno_ibov - retorno_cdi,
            "vencedor": vencedor
        })

        datas_usadas.append(data_entrada)

    ciclos_df = pd.DataFrame(ciclos)

    if ciclos_df.empty:
        return ciclos_df

    ciclos_df = ciclos_df.sort_values("diferenca", ascending=False)
    ciclos_df = ciclos_df.head(numero_maximo_ciclos)

    return ciclos_df

# ============================================================
# 3.1 JANELAS ESTRATÉGICAS OBRIGATÓRIAS
# ============================================================

def adicionar_janelas_estrategicas(df: pd.DataFrame, ciclos_df: pd.DataFrame) -> pd.DataFrame:
    """
    Adiciona janelas estratégicas relevantes para leitura consultiva,
    mesmo quando elas não aparecem entre os maiores ciclos matemáticos.

    Janela inicial: Covid-19.
    """

    if df.empty:
        return ciclos_df

    base = df.copy()
    base = base.sort_values("data")

    janelas = []

    inicio_covid = pd.Timestamp("2020-02-01")
    fim_covid_busca = pd.Timestamp("2020-04-30")

    periodo_contem_covid = (
        base["data"].min() <= inicio_covid
        and base["data"].max() >= fim_covid_busca
    )

    if periodo_contem_covid:
        janela_fundo = base[
            (base["data"] >= inicio_covid)
            & (base["data"] <= fim_covid_busca)
        ].copy()

        if not janela_fundo.empty:
            fundo = janela_fundo.loc[janela_fundo["ibovespa"].idxmin()]

            data_entrada = fundo["data"]
            ibov_entrada = fundo["ibovespa"]

            saida = base.iloc[-1]
            data_saida = saida["data"]
            ibov_saida = saida["ibovespa"]

            janela = base[
                (base["data"] >= data_entrada)
                & (base["data"] <= data_saida)
            ].copy()

            if not janela.empty:
                retorno_ibov = ibov_saida / ibov_entrada - 1
                retorno_cdi = janela["indice_cdi"].iloc[-1] / janela["indice_cdi"].iloc[0] - 1

                vencedor = "Ibovespa" if retorno_ibov > retorno_cdi else "CDI"

                janelas.append(
                    {
                        "entrada": data_entrada.date(),
                        "saida": data_saida.date(),
                        "ibov_entrada": ibov_entrada,
                        "ibov_saida": ibov_saida,
                        "retorno_ibov": retorno_ibov,
                        "retorno_cdi": retorno_cdi,
                        "diferenca": retorno_ibov - retorno_cdi,
                        "vencedor": vencedor,
                        "tipo": "Janela estratégica",
                        "contexto": "Covid-19, circuit breakers, choque global de risco e recuperação posterior",
                        "leitura": (
                            "Janela relevante para demonstrar como entradas em momentos de estresse "
                            "podem alterar fortemente o retorno relativo do Ibovespa frente ao CDI, "
                            "desde que exista horizonte, tolerância à volatilidade e aderência ao perfil."
                        ),
                    }
                )

    if ciclos_df is None or ciclos_df.empty:
        if janelas:
            return pd.DataFrame(janelas)
        return pd.DataFrame()

    ciclos_base = ciclos_df.copy()

    if "tipo" not in ciclos_base.columns:
        ciclos_base["tipo"] = "Ciclo detectado"

    if "contexto" not in ciclos_base.columns:
        ciclos_base["contexto"] = ""

    if "leitura" not in ciclos_base.columns:
        ciclos_base["leitura"] = ""

    if not janelas:
        return ciclos_base

    janelas_df = pd.DataFrame(janelas)

    ciclos_final = pd.concat(
        [ciclos_base, janelas_df],
        ignore_index=True
    )

    ciclos_final = ciclos_final.drop_duplicates(
        subset=["entrada", "saida"],
        keep="last"
    )

    return ciclos_final


# ============================================================
# 4. GRÁFICOS
# ============================================================

def grafico_ibovespa_cdi(df: pd.DataFrame) -> go.Figure:
    base = df.copy()

    valor_final_ibov = base["indice_ibov"].iloc[-1]
    valor_final_cdi = base["indice_cdi"].iloc[-1]

    fig = go.Figure()

    fig.add_trace(
        go.Scatter(
            x=base["data"],
            y=base["indice_ibov"],
            mode="lines",
            name="Ibovespa",
            line=dict(color="#00D4FF", width=3),
            hovertemplate=(
                "<b>Ibovespa</b><br>"
                "Data: %{x|%d/%m/%Y}<br>"
                "Índice base 100: %{y:.2f}<extra></extra>"
            ),
        )
    )

    fig.add_trace(
        go.Scatter(
            x=base["data"],
            y=base["indice_cdi"],
            mode="lines",
            name="CDI",
            line=dict(color="#00E676", width=3, dash="solid"),
            hovertemplate=(
                "<b>CDI</b><br>"
                "Data: %{x|%d/%m/%Y}<br>"
                "Índice base 100: %{y:.2f}<extra></extra>"
            ),
            fill="tozeroy",
            fillcolor="rgba(0, 230, 118, 0.06)"
        )
    )

    fig.add_annotation(
        x=base["data"].iloc[-1],
        y=valor_final_ibov,
        text=f"Ibovespa: {valor_final_ibov:.1f}",
        showarrow=True,
        arrowhead=2,
        ax=40,
        ay=-40,
        font=dict(size=12, color="#00D4FF"),
        bgcolor="rgba(0,0,0,0.45)",
        bordercolor="#00D4FF",
        borderwidth=1
    )

    fig.add_annotation(
        x=base["data"].iloc[-1],
        y=valor_final_cdi,
        text=f"CDI: {valor_final_cdi:.1f}",
        showarrow=True,
        arrowhead=2,
        ax=40,
        ay=40,
        font=dict(size=12, color="#00E676"),
        bgcolor="rgba(0,0,0,0.45)",
        bordercolor="#00E676",
        borderwidth=1
    )

    eventos_historicos = [
        {
            "data": "2002-10-01",
            "titulo": "2002",
            "texto": "Risco Brasil, câmbio e eleição presidencial"
        },
        {
            "data": "2008-09-15",
            "titulo": "2008",
            "texto": "Crise financeira global"
        },
        {
            "data": "2016-08-31",
            "titulo": "2015/16",
            "texto": "Recessão brasileira e crise política/fiscal"
        },
        {
            "data": "2020-03-23",
            "titulo": "Covid",
            "texto": "Pandemia, circuit breakers e choque global de risco"
        },
        {
            "data": "2022-03-01",
            "titulo": "2021/22",
            "texto": "Inflação, alta da Selic e guerra na Ucrânia"
        },
        {
            "data": "2024-01-01",
            "titulo": "2023/26",
            "texto": "Juros, fiscal, commodities e fluxo estrangeiro"
        },
    ]

    data_min = base["data"].min()
    data_max = base["data"].max()
    y_ref = max(valor_final_ibov, valor_final_cdi)

    for evento in eventos_historicos:
        data_evento = pd.Timestamp(evento["data"])

        if data_min <= data_evento <= data_max:
            fig.add_vline(
                x=data_evento,
                line_width=1,
                line_dash="dot",
                line_color="rgba(255,255,255,0.35)"
            )

            fig.add_annotation(
                x=data_evento,
                y=y_ref,
                text=f"{evento['titulo']}<br>{evento['texto']}",
                showarrow=False,
                yshift=45,
                font=dict(size=10, color="#EAEAEA"),
                bgcolor="rgba(0,0,0,0.55)",
                bordercolor="rgba(255,255,255,0.25)",
                borderwidth=1
            )

    fig.update_layout(
        title={
            "text": "Comparativo histórico | Ibovespa x CDI",
            "x": 0.02,
            "xanchor": "left"
        },
        height=560,
        hovermode="x unified",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#EAEAEA"),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        ),
        margin=dict(l=20, r=20, t=70, b=20),
        xaxis=dict(
            title="Data",
            showgrid=False,
            rangeslider=dict(visible=True),
            tickfont=dict(color="#CFCFCF")
        ),
        yaxis=dict(
            title="Índice base 100",
            showgrid=True,
            gridcolor="rgba(255,255,255,0.08)",
            zeroline=False,
            tickfont=dict(color="#CFCFCF")
        )
    )

    return fig

def grafico_retorno_percentual(df: pd.DataFrame) -> go.Figure:
    base = df.copy()

    base["retorno_pct_ibov"] = base["retorno_acum_ibov"] * 100
    base["retorno_pct_cdi"] = base["retorno_acum_cdi"] * 100

    fig = go.Figure()

    fig.add_trace(
        go.Scatter(
            x=base["data"],
            y=base["retorno_pct_ibov"],
            mode="lines",
            name="Ibovespa (%)",
            line=dict(color="#42A5F5", width=3),
            hovertemplate=(
                "<b>Ibovespa</b><br>"
                "Data: %{x|%d/%m/%Y}<br>"
                "Retorno acumulado: %{y:.2f}%<extra></extra>"
            ),
        )
    )

    fig.add_trace(
        go.Scatter(
            x=base["data"],
            y=base["retorno_pct_cdi"],
            mode="lines",
            name="CDI (%)",
            line=dict(color="#66BB6A", width=3),
            hovertemplate=(
                "<b>CDI</b><br>"
                "Data: %{x|%d/%m/%Y}<br>"
                "Retorno acumulado: %{y:.2f}%<extra></extra>"
            ),
        )
    )

    fig.add_hline(
        y=0,
        line_width=1,
        line_dash="dot",
        line_color="rgba(255,255,255,0.25)"
    )

    fig.update_layout(
        title={
            "text": "Retorno acumulado (%) | Ibovespa x CDI",
            "x": 0.02,
            "xanchor": "left"
        },
        height=480,
        hovermode="x unified",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#EAEAEA"),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        ),
        margin=dict(l=20, r=20, t=70, b=20),
        xaxis=dict(
            title="Data",
            showgrid=False
        ),
        yaxis=dict(
            title="Retorno acumulado (%)",
            showgrid=True,
            gridcolor="rgba(255,255,255,0.08)",
            zeroline=False
        )
    )

    return fig


def grafico_diferenca(df: pd.DataFrame) -> go.Figure:
    base = df.copy()
    base["diferenca_ibov_cdi"] = base["indice_ibov"] - base["indice_cdi"]

    fig = go.Figure()

    fig.add_trace(
        go.Scatter(
            x=base["data"],
            y=base["diferenca_ibov_cdi"],
            mode="lines",
            name="Diferença Ibovespa - CDI",
            hovertemplate="Data: %{x}<br>Diferença: %{y:.2f}<extra></extra>"
        )
    )

    fig.update_layout(
        title="Diferença acumulada entre Ibovespa e CDI",
        xaxis_title="Data",
        yaxis_title="Diferença em pontos do índice base 100",
        hovermode="x unified",
        height=420
    )

    return fig


# ============================================================
# 5. RELATÓRIO FORESIGHT
# ============================================================

def gerar_texto_foresight(metricas: dict, ciclos_df: pd.DataFrame) -> str:
    """
    Gera uma leitura consultiva inicial.
    Depois podemos evoluir isso para LLM/API.
    """

    if not metricas:
        return "Não foi possível gerar o relatório por ausência de dados suficientes."

    texto = f"""
Relatório Foresight | Ibovespa x CDI

O módulo compara o desempenho histórico do Ibovespa com o CDI, considerando uma leitura de retorno, risco, volatilidade e comportamento de ciclos. No período analisado, entre {metricas["data_inicio"].date()} e {metricas["data_fim"].date()}, o Ibovespa apresentou retorno acumulado de {formatar_pct(metricas["retorno_ibov"])}, enquanto o CDI acumulou {formatar_pct(metricas["retorno_cdi"])}. Em termos anualizados, o Ibovespa apresentou retorno médio de {formatar_pct(metricas["retorno_anual_ibov"])} ao ano, contra {formatar_pct(metricas["retorno_anual_cdi"])} ao ano do CDI. O vencedor do período foi: {metricas["vencedor"]}.

A leitura do Ibovespa exige atenção a múltiplos fatores. A bolsa brasileira tende a reagir à trajetória da taxa de juros, inflação, câmbio, fluxo de investidores estrangeiros, desempenho de commodities, resultados de empresas com grande peso no índice, cenário fiscal, política monetária internacional e eventos de aversão ou retomada de risco. Petrobras, Vale, bancos e empresas ligadas ao ciclo econômico costumam exercer influência relevante sobre o índice, o que torna a análise setorial indispensável para interpretar movimentos mais fortes.

Do ponto de vista comportamental, o Ibovespa também é afetado por excesso de pessimismo, euforia, comportamento de manada, aversão à perda e busca por liquidez em momentos de estresse. Grandes quedas costumam concentrar medo, incerteza e saída de investidores, enquanto ciclos de recuperação tendem a ocorrer quando há melhora gradual de expectativas, queda projetada de juros, recomposição de lucros corporativos ou redução da percepção de risco.

O CDI, por sua vez, é diretamente influenciado pela política monetária. Sua dinâmica acompanha a taxa básica de juros, as decisões do Copom, as expectativas de inflação, o prêmio de risco da economia e a necessidade de aperto ou afrouxamento monetário. Por isso, o CDI tende a oferecer maior previsibilidade, menor volatilidade e trajetória mais estável, embora possa perder atratividade relativa em ciclos de queda de juros.

A comparação histórica entre Ibovespa e CDI não deve ser lida como uma escolha binária entre risco e segurança. O CDI cumpre papel de estabilidade, liquidez e previsibilidade. O Ibovespa, por outro lado, pode entregar prêmio superior em janelas específicas, especialmente quando o ponto de entrada ocorre após quedas relevantes e quando o investidor possui horizonte suficiente para atravessar volatilidade.


"""

    if ciclos_df is not None and not ciclos_df.empty:
        texto += """

Ciclos macro para leitura consultiva

Os 8 ciclos macro mais úteis para leitura consultiva seriam:

• 1968–1971 — ciclo inicial de forte valorização da bolsa brasileira.
• 1971–1975 — estouro da bolha e forte correção.
• 1980s/início dos 1990s — inflação alta, planos econômicos e instabilidade monetária.
• 1994–1997 — Plano Real, estabilização e reprecificação de ativos.
• 1998–2002 — crises externa, câmbio, risco Brasil e eleição de 2002.
• 2003–2008 — boom de commodities, China, crédito e forte ciclo de alta.
• 2008–2016 — crise financeira global, recuperação parcial, recessão brasileira e crise política/fiscal.
• 2020–2026 — Covid, Selic baixa, inflação, alta de juros, commodities, fiscal e recuperação posterior.

"""

        for _, row in ciclos_df.iterrows():
            texto += (
                f"- Entre {row['entrada']} e {row['saida']}, o Ibovespa saiu de "
                f"{formatar_numero(row['ibov_entrada'])} pontos para "
                f"{formatar_numero(row['ibov_saida'])} pontos. "
                f"No período, o Ibovespa acumulou {formatar_pct(row['retorno_ibov'])}, "
                f"contra {formatar_pct(row['retorno_cdi'])} do CDI. "
                f"O vencedor da janela foi: {row['vencedor']}.\n"
            )

        texto += "\n"

    texto += f"""
A leitura estratégica final indica que o Ibovespa deve ser observado com maior atenção em momentos de queda expressiva, juros em perspectiva de redução, melhora de expectativas econômicas, entrada de fluxo estrangeiro, valorização de commodities ou recuperação de resultados corporativos. Esses sinais não eliminam o risco, mas ajudam a qualificar a conversa com o cliente.

Para o cliente conservador, o CDI permanece como referência de segurança e previsibilidade. Para o cliente com maior tolerância a risco, horizonte alongado e capacidade de suportar volatilidade, o Ibovespa pode compor uma parcela estratégica da carteira, especialmente quando a entrada ocorre em ciclos de assimetria favorável. A decisão, contudo, deve sempre considerar perfil, liquidez, horizonte, objetivos patrimoniais e adequação regulatória.
"""

    return texto.strip()

def fig_to_png_bytes(fig: go.Figure) -> BytesIO:
    """
    Converte um gráfico Plotly em imagem PNG para inserir no Word.
    Requer kaleido instalado.
    """

    image_bytes = fig.to_image(
        format="png",
        width=1200,
        height=650,
        scale=2
    )

    buffer = BytesIO(image_bytes)
    buffer.seek(0)

    return buffer

    # ============================================================
# 6. INTERFACE STREAMLIT DO MÓDULO
# ============================================================

def render_ibovespa_cdi_module():
    st.markdown("### Painel analítico Ibovespa x CDI")

    st.caption(
        "Comparação histórica entre bolsa brasileira e CDI, com leitura de ciclos, risco e relatório consultivo."
    )

    col_periodo_1, col_periodo_2, col_parametros = st.columns([1, 1, 1])

    with col_periodo_1:
        data_inicio = st.date_input(
            "Data inicial",
            value=date(2000, 1, 1),
            key="ibov_cdi_data_inicio"
        )

    with col_periodo_2:
        data_fim = st.date_input(
            "Data final",
            value=date.today(),
            key="ibov_cdi_data_fim"
        )

    with col_parametros:
        numero_ciclos = st.slider(
            "Máximo de ciclos relevantes",
            min_value=1,
            max_value=15,
            value=8,
            key="ibov_cdi_numero_ciclos"
        )

    with st.expander("Parâmetros técnicos da detecção de ciclos", expanded=False):
        col_a, col_b, col_c = st.columns(3)

        with col_a:
            queda_minima = st.slider(
                "Queda mínima para buscar fundo",
                min_value=5,
                max_value=60,
                value=20,
                step=5,
                format="%d%%",
                key="ibov_cdi_queda_minima"
            ) / 100

        with col_b:
            alta_minima = st.slider(
                "Alta mínima após entrada",
                min_value=10,
                max_value=150,
                value=35,
                step=5,
                format="%d%%",
                key="ibov_cdi_alta_minima"
            ) / 100

        with col_c:
            meses_minimos = st.slider(
                "Janela mínima em meses",
                min_value=3,
                max_value=60,
                value=12,
                step=3,
                key="ibov_cdi_meses_minimos"
            )

    if data_inicio >= data_fim:
        st.warning("A data inicial precisa ser anterior à data final.")
        return

    with st.spinner("Carregando séries históricas do Ibovespa e do CDI..."):
        df_ibov = carregar_ibovespa(str(data_inicio), str(data_fim))
        df_cdi = carregar_cdi_bacen(str(data_inicio), str(data_fim))

    if df_ibov.empty:
        st.error("Não foi possível carregar o histórico do Ibovespa.")
        return

    if df_cdi.empty:
        st.error("Não foi possível carregar o histórico do CDI pelo Bacen.")
        return

    df = preparar_base_comparativa(df_ibov, df_cdi)

    if df.empty:
        st.error("Não foi possível montar a base comparativa Ibovespa x CDI.")
        return

    metricas = calcular_metricas(df)

    ciclos_df = detectar_ciclos_relevantes(
        df=df,
        numero_maximo_ciclos=numero_ciclos,
        queda_minima=queda_minima,
        alta_minima=alta_minima,
        meses_minimos=meses_minimos
    )

    ciclos_df = adicionar_janelas_estrategicas(df, ciclos_df)

    # ============================================================
    # PAINEL DE MÉTRICAS
    # ============================================================

    st.markdown("### Painel comparativo")

    col1, col2, col3, col4 = st.columns(4)

    col1.metric(
        "Ibovespa acumulado",
        formatar_pct(metricas["retorno_ibov"])
    )

    col2.metric(
        "CDI acumulado",
        formatar_pct(metricas["retorno_cdi"])
    )

    col3.metric(
        "Vencedor histórico",
        metricas["vencedor"]
    )

    col4.metric(
        "Volatilidade Ibovespa",
        formatar_pct(metricas["volatilidade_ibov"])
    )

    col5, col6, col7, col8 = st.columns(4)

    col5.metric(
        "Ibovespa ao ano",
        formatar_pct(metricas["retorno_anual_ibov"])
    )

    col6.metric(
        "CDI ao ano",
        formatar_pct(metricas["retorno_anual_cdi"])
    )

    col7.metric(
        "Ibovespa ao mês",
        formatar_pct(metricas["retorno_mensal_ibov"])
    )

    col8.metric(
        "Drawdown Ibovespa",
        formatar_pct(metricas["max_drawdown_ibov"])
    )

    # ============================================================
    # GRÁFICOS
    # ============================================================

    st.markdown("### Evolução comparativa")

    st.info(
        "Leitura visual comparativa entre Ibovespa e CDI. "
        "O gráfico base 100 mostra como R$ 100 evoluiriam em cada alternativa ao longo do tempo."
    )

    st.plotly_chart(
        grafico_ibovespa_cdi(df),
        use_container_width=True
    )

    st.plotly_chart(
        grafico_retorno_percentual(df),
        use_container_width=True
    )

    st.plotly_chart(
        grafico_diferenca(df),
        use_container_width=True
    )

    # ============================================================
    # CICLOS RELEVANTES
    # ============================================================

    st.markdown("### Ciclos relevantes e janelas estratégicas")

    if ciclos_df.empty:
        st.info(
            "Nenhum ciclo relevante foi identificado com os parâmetros atuais. "
            "Você pode reduzir a queda mínima, reduzir a alta mínima ou ampliar o período analisado."
        )
    else:
        ciclos_view = ciclos_df.copy()

        if "tipo" not in ciclos_view.columns:
            ciclos_view["tipo"] = "Ciclo detectado"

        if "contexto" not in ciclos_view.columns:
            ciclos_view["contexto"] = ""

        if "leitura" not in ciclos_view.columns:
            ciclos_view["leitura"] = ""

        ciclos_view["ibov_entrada"] = ciclos_view["ibov_entrada"].apply(formatar_numero)
        ciclos_view["ibov_saida"] = ciclos_view["ibov_saida"].apply(formatar_numero)
        ciclos_view["retorno_ibov"] = ciclos_view["retorno_ibov"].apply(formatar_pct)
        ciclos_view["retorno_cdi"] = ciclos_view["retorno_cdi"].apply(formatar_pct)
        ciclos_view["diferenca"] = ciclos_view["diferenca"].apply(formatar_pct)

        ciclos_view = ciclos_view.rename(
            columns={
                "entrada": "Entrada",
                "saida": "Saída",
                "ibov_entrada": "Ibovespa na entrada",
                "ibov_saida": "Ibovespa na saída",
                "retorno_ibov": "Retorno Ibovespa",
                "retorno_cdi": "Retorno CDI",
                "diferenca": "Diferença",
                "vencedor": "Vencedor",
                "tipo": "Tipo",
                "contexto": "Contexto",
                "leitura": "Leitura consultiva",
            }
        )

        colunas_exibir = [
            "Tipo",
            "Entrada",
            "Saída",
            "Ibovespa na entrada",
            "Ibovespa na saída",
            "Retorno Ibovespa",
            "Retorno CDI",
            "Diferença",
            "Vencedor",
            "Contexto",
            "Leitura consultiva",
        ]

        colunas_exibir = [
            coluna for coluna in colunas_exibir
            if coluna in ciclos_view.columns
        ]

        st.dataframe(
            ciclos_view[colunas_exibir],
            use_container_width=True,
            hide_index=True
        )

    # ============================================================
    # CICLOS MACRO
    # ============================================================

    st.markdown("### Ciclos macro para leitura consultiva")

    st.info(
        "Além dos ciclos detectados pelo algoritmo, o módulo considera oito grandes ciclos "
        "macroeconômicos do Ibovespa que ajudam a contextualizar a conversa consultiva."
    )

    ciclos_macro_df = pd.DataFrame(
        [
            {
                "Ciclo": "1968–1971",
                "Contexto": "Ciclo inicial de forte valorização da bolsa brasileira",
                "Leitura consultiva": "Fase de formação e expansão inicial do mercado acionário brasileiro."
            },
            {
                "Ciclo": "1971–1975",
                "Contexto": "Estouro da bolha e forte correção",
                "Leitura consultiva": "Mostra que ciclos de euforia podem ser seguidos por ajustes profundos."
            },
            {
                "Ciclo": "1980s/início dos 1990s",
                "Contexto": "Inflação alta, planos econômicos e instabilidade monetária",
                "Leitura consultiva": "Período marcado por distorções monetárias, perda de referência de preços e elevada incerteza."
            },
            {
                "Ciclo": "1994–1997",
                "Contexto": "Plano Real, estabilização e reprecificação de ativos",
                "Leitura consultiva": "A estabilização monetária alterou expectativas e permitiu nova leitura de valor dos ativos."
            },
            {
                "Ciclo": "1998–2002",
                "Contexto": "Crises externa, câmbio, risco Brasil e eleição de 2002",
                "Leitura consultiva": "Janela importante para observar prêmio de risco, câmbio e incerteza política."
            },
            {
                "Ciclo": "2003–2008",
                "Contexto": "Boom de commodities, China, crédito e forte ciclo de alta",
                "Leitura consultiva": "Período de valorização associado à liquidez global, crescimento chinês e empresas de commodities."
            },
            {
                "Ciclo": "2008–2016",
                "Contexto": "Crise financeira global, recuperação parcial, recessão brasileira e crise política/fiscal",
                "Leitura consultiva": "Ciclo longo de alternância entre recuperação, deterioração doméstica e aumento da percepção de risco."
            },
            {
                "Ciclo": "2020–2026",
                "Contexto": "Covid, Selic baixa, inflação, alta de juros, commodities, fiscal e recuperação posterior",
                "Leitura consultiva": "Janela essencial para discutir entrada em momentos de estresse, horizonte longo e recuperação pós-choque."
            },
        ]
    )

    st.dataframe(
        ciclos_macro_df,
        use_container_width=True,
        hide_index=True
    )

    # ============================================================
    # RELATÓRIO FORESIGHT
    # ============================================================

    st.markdown("### Relatório Foresight")

    relatorio = gerar_texto_foresight(metricas, ciclos_df)

    with st.expander("Visualizar relatório consultivo", expanded=True):
        st.write(relatorio)

    arquivo_word = gerar_word_relatorio(
        texto=relatorio,
        metricas=metricas,
        ciclos_df=ciclos_df,
        df=df
    )

    st.download_button(
        label="Baixar relatório Ibovespa x CDI em Word",
        data=arquivo_word,
        file_name="relatorio_foresight_ibovespa_cdi.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    with st.expander("Observação metodológica", expanded=False):
        st.write(
            """
            Este módulo tem finalidade analítica e consultiva. 
            A comparação histórica entre Ibovespa e CDI não constitui recomendação individualizada de investimento. 
            A interpretação deve considerar perfil do investidor, horizonte, liquidez, adequação regulatória, tolerância a risco e composição global da carteira.
            """
        )

def gerar_word_relatorio(
    texto: str,
    metricas: dict,
    ciclos_df: pd.DataFrame,
    df: pd.DataFrame
) -> BytesIO:

    doc = Document()

    doc.add_heading("Relatório Foresight | Ibovespa x CDI", level=1)

    doc.add_paragraph(
        "Relatório consultivo gerado a partir da comparação histórica entre "
        "Ibovespa e CDI, com leitura de retorno, risco, ciclos de mercado "
        "e interpretação estratégica."
    )

    # ============================================================
    # RESUMO EXECUTIVO
    # ============================================================

    doc.add_heading("Resumo Executivo", level=2)

    for paragrafo in texto.split("\n\n"):
        if paragrafo.strip():
            doc.add_paragraph(paragrafo.strip())

    # ============================================================
    # MÉTRICAS COMPARATIVAS
    # ============================================================

    doc.add_heading("Métricas Comparativas", level=2)

    tabela = doc.add_table(rows=1, cols=3)
    tabela.style = "Table Grid"

    hdr = tabela.rows[0].cells
    hdr[0].text = "Métrica"
    hdr[1].text = "Ibovespa"
    hdr[2].text = "CDI"

    linhas = [
        (
            "Retorno acumulado",
            formatar_pct(metricas["retorno_ibov"]),
            formatar_pct(metricas["retorno_cdi"])
        ),
        (
            "Retorno anualizado",
            formatar_pct(metricas["retorno_anual_ibov"]),
            formatar_pct(metricas["retorno_anual_cdi"])
        ),
        (
            "Retorno mensal médio",
            formatar_pct(metricas["retorno_mensal_ibov"]),
            formatar_pct(metricas["retorno_mensal_cdi"])
        ),
        (
            "Volatilidade anualizada",
            formatar_pct(metricas["volatilidade_ibov"]),
            "Não aplicável no modelo simplificado"
        ),
        (
            "Maior drawdown",
            formatar_pct(metricas["max_drawdown_ibov"]),
            "Não aplicável no modelo simplificado"
        ),
        (
            "Vencedor no período",
            metricas["vencedor"],
            metricas["vencedor"]
        ),
    ]

    for metrica, ibov, cdi in linhas:
        cells = tabela.add_row().cells
        cells[0].text = metrica
        cells[1].text = ibov
        cells[2].text = cdi

    # ============================================================
    # CICLOS RELEVANTES E JANELAS ESTRATÉGICAS
    # ============================================================

    if ciclos_df is not None and not ciclos_df.empty:
        doc.add_heading("Ciclos relevantes e janelas estratégicas", level=2)

        tabela_ciclos = doc.add_table(rows=1, cols=9)
        tabela_ciclos.style = "Table Grid"

        hdr = tabela_ciclos.rows[0].cells
        hdr[0].text = "Tipo"
        hdr[1].text = "Entrada"
        hdr[2].text = "Saída"
        hdr[3].text = "Ibovespa entrada"
        hdr[4].text = "Ibovespa saída"
        hdr[5].text = "Retorno Ibovespa"
        hdr[6].text = "Retorno CDI"
        hdr[7].text = "Vencedor"
        hdr[8].text = "Contexto"

        ciclos_word = ciclos_df.copy()

        if "tipo" not in ciclos_word.columns:
            ciclos_word["tipo"] = "Ciclo detectado"

        if "contexto" not in ciclos_word.columns:
            ciclos_word["contexto"] = ""

        for _, row in ciclos_word.iterrows():
            cells = tabela_ciclos.add_row().cells
            cells[0].text = str(row.get("tipo", "Ciclo detectado"))
            cells[1].text = str(row["entrada"])
            cells[2].text = str(row["saida"])
            cells[3].text = formatar_numero(row["ibov_entrada"])
            cells[4].text = formatar_numero(row["ibov_saida"])
            cells[5].text = formatar_pct(row["retorno_ibov"])
            cells[6].text = formatar_pct(row["retorno_cdi"])
            cells[7].text = str(row["vencedor"])
            cells[8].text = str(row.get("contexto", ""))

    # ============================================================
    # CICLOS MACRO PARA LEITURA CONSULTIVA
    # ============================================================

    doc.add_heading("Ciclos macro para leitura consultiva", level=2)

    doc.add_paragraph(
        "Os 8 ciclos macro mais úteis para leitura consultiva seriam:"
    )

    ciclos_macro_word = [
        ("1968–1971", "Ciclo inicial de forte valorização da bolsa brasileira."),
        ("1971–1975", "Estouro da bolha e forte correção."),
        ("1980s/início dos 1990s", "Inflação alta, planos econômicos e instabilidade monetária."),
        ("1994–1997", "Plano Real, estabilização e reprecificação de ativos."),
        ("1998–2002", "Crises externa, câmbio, risco Brasil e eleição de 2002."),
        ("2003–2008", "Boom de commodities, China, crédito e forte ciclo de alta."),
        ("2008–2016", "Crise financeira global, recuperação parcial, recessão brasileira e crise política/fiscal."),
        ("2020–2026", "Covid, Selic baixa, inflação, alta de juros, commodities, fiscal e recuperação posterior."),
    ]

    tabela_macro = doc.add_table(rows=1, cols=2)
    tabela_macro.style = "Table Grid"

    hdr = tabela_macro.rows[0].cells
    hdr[0].text = "Ciclo"
    hdr[1].text = "Leitura macroeconômica"

    for ciclo, leitura in ciclos_macro_word:
        cells = tabela_macro.add_row().cells
        cells[0].text = ciclo
        cells[1].text = leitura

    # ============================================================
    # GRÁFICOS DO RADAR
    # ============================================================

    doc.add_heading("Gráficos comparativos", level=2)

    try:
        fig_base_100 = grafico_ibovespa_cdi(df)
        imagem_base_100 = fig_to_png_bytes(fig_base_100)

        doc.add_paragraph(
            "Evolução comparativa em base 100: mostra como R$ 100 evoluiriam "
            "em cada alternativa ao longo do período analisado."
        )

        doc.add_picture(imagem_base_100, width=Inches(6.5))

    except Exception as error:
        doc.add_paragraph(
            f"Não foi possível inserir o gráfico Ibovespa x CDI no relatório. Erro: {error}"
        )

    try:
        fig_retorno = grafico_retorno_percentual(df)
        imagem_retorno = fig_to_png_bytes(fig_retorno)

        doc.add_paragraph(
            "Retorno acumulado percentual: compara o desempenho acumulado do Ibovespa "
            "e do CDI ao longo do período."
        )

        doc.add_picture(imagem_retorno, width=Inches(6.5))

    except Exception as error:
        doc.add_paragraph(
            f"Não foi possível inserir o gráfico de retorno percentual no relatório. Erro: {error}"
        )

    try:
        fig_diferenca = grafico_diferenca(df)
        imagem_diferenca = fig_to_png_bytes(fig_diferenca)

        doc.add_paragraph(
            "Diferença acumulada entre Ibovespa e CDI: evidencia os períodos em que "
            "a bolsa ampliou ou reduziu vantagem relativa frente ao CDI."
        )

        doc.add_picture(imagem_diferenca, width=Inches(6.5))

    except Exception as error:
        doc.add_paragraph(
            f"Não foi possível inserir o gráfico de diferença no relatório. Erro: {error}"
        )

    # ============================================================
    # AVISO METODOLÓGICO
    # ============================================================

    doc.add_heading("Aviso metodológico", level=2)

    doc.add_paragraph(
        "Esta análise tem finalidade exclusivamente educacional, analítica e consultiva. "
        "A comparação histórica entre Ibovespa e CDI não constitui recomendação individualizada "
        "de investimento, promessa de rentabilidade ou garantia de desempenho futuro. "
        "A decisão de investimento deve considerar perfil do cliente, liquidez, horizonte, "
        "objetivos patrimoniais, tolerância a risco e adequação regulatória."
    )

    # ============================================================
    # SALVAR DOCUMENTO
    # ============================================================

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer