from io import BytesIO
import re

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================================================
# IDENTIDADE VISUAL DO RELATÓRIO
# =========================================================

COLOR_NAVY = RGBColor(11, 31, 58)
COLOR_BLUE = RGBColor(0, 92, 169)
COLOR_WHITE = RGBColor(255, 255, 255)
COLOR_DARK_GRAY = RGBColor(51, 51, 51)
COLOR_MUTED_GRAY = RGBColor(100, 100, 100)

COLOR_NAVY_HEX = "0B1F3A"
COLOR_LIGHT_BLUE_HEX = "DCEEFF"
COLOR_GRAY_HEX = "F4F6F8"
COLOR_BORDER_HEX = "D9DEE5"
COLOR_CARD_BORDER_HEX = "BFD7F2"

FONT_TITLE = "Arial"
FONT_BODY = "Arial"
FONT_TABLE = "Arial"


# =========================================================
# FUNÇÕES DE TEXTO
# =========================================================

def clean_markdown(text):
    if text is None:
        return ""

    text = str(text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("###", "")
    text = text.replace("##", "")
    text = text.replace("#", "")

    return text.strip()


def format_currency(value):
    try:
        value = float(value)
        formatted = f"R$ {value:,.2f}"
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
        return formatted
    except Exception:
        return "" if value is None else str(value)


def format_percent(value):
    try:
        value = float(value)
        return f"{value:.2f}%".replace(".", ",")
    except Exception:
        return "" if value is None else str(value)


def format_number(value):
    try:
        value = float(value)
        return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "" if value is None else str(value)


def normalize_brl_currency_and_percent_in_text(text):
    """
    Corrige valores e percentuais que venham em padrão americano dentro dos textos:
    R$ 739,174.15 -> R$ 739.174,15
    12.74% -> 12,74%
    """

    if text is None:
        return ""

    text = str(text)

    currency_pattern = r"R\$\s?(\d{1,3}(?:,\d{3})+\.\d{2})"

    def replace_currency(match):
        raw_value = match.group(1)

        try:
            numeric_value = float(raw_value.replace(",", ""))
            return format_currency(numeric_value)
        except Exception:
            return match.group(0)

    text = re.sub(currency_pattern, replace_currency, text)

    percent_pattern = r"(?<!\d)(\d+\.\d{1,4})%"

    def replace_percent(match):
        raw_value = match.group(1)

        try:
            numeric_value = float(raw_value)
            return format_percent(numeric_value)
        except Exception:
            return match.group(0)

    text = re.sub(percent_pattern, replace_percent, text)

    return text


def safe_text(value):
    if value is None:
        return ""

    text = clean_markdown(value)
    text = normalize_brl_currency_and_percent_in_text(text)

    return str(text)


# =========================================================
# FUNÇÕES DE FORMATAÇÃO DO WORD
# =========================================================

def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(
    cell,
    color=COLOR_BORDER_HEX,
    size="4",
    space="0"
):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_borders = tc_pr.first_child_found_in("w:tcBorders")

    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), space)
        element.set(qn("w:color"), color)


def set_cell_text(
    cell,
    text,
    bold=False,
    font_size=8,
    color=None,
    align="center"
):
    cell.text = ""

    paragraph = cell.paragraphs[0]

    if align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(safe_text(text))
    run.bold = bold
    run.font.name = FONT_TABLE
    run.font.size = Pt(font_size)

    if color is not None:
        run.font.color.rgb = color

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(document, text, bold=False, font_size=10):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run(safe_text(text))
    run.bold = bold
    run.font.name = FONT_BODY
    run.font.size = Pt(font_size)
    run.font.color.rgb = COLOR_DARK_GRAY

    return paragraph


def add_section_heading(document, number, title):
    paragraph = document.add_paragraph()
    paragraph.style = "Heading 1"

    run = paragraph.add_run(f"{number}. {title}")
    run.bold = True
    run.font.name = FONT_TITLE
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_NAVY

    return paragraph


def add_subheading(document, title):
    paragraph = document.add_paragraph()
    paragraph.style = "Heading 2"

    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_BLUE

    return paragraph


def add_key_value_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for label, value in rows:
        row_cells = table.add_row().cells

        shade_cell(row_cells[0], COLOR_GRAY_HEX)
        set_cell_border(row_cells[0])
        set_cell_border(row_cells[1])

        set_cell_text(
            row_cells[0],
            label,
            bold=True,
            font_size=8,
            color=COLOR_NAVY,
            align="left"
        )

        set_cell_text(
            row_cells[1],
            value,
            font_size=8,
            color=COLOR_DARK_GRAY,
            align="left"
        )

    document.add_paragraph()
    return table


def add_dataframe_table(document, df, max_rows=None, font_size=7):
    if df is None or df.empty:
        add_paragraph(document, "Não há dados disponíveis para esta seção.")
        return None

    table_df = df.copy()

    if max_rows is not None:
        table_df = table_df.head(max_rows)

    table = document.add_table(rows=1, cols=len(table_df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for index, column in enumerate(table_df.columns):
        cell = header_cells[index]

        shade_cell(cell, COLOR_NAVY_HEX)
        set_cell_border(cell, color=COLOR_NAVY_HEX, size="4")

        set_cell_text(
            cell,
            str(column),
            bold=True,
            font_size=font_size,
            color=COLOR_WHITE
        )

    for row_index, (_, row) in enumerate(table_df.iterrows()):
        row_cells = table.add_row().cells

        for index, value in enumerate(row):
            cell = row_cells[index]

            if row_index % 2 == 0:
                shade_cell(cell, COLOR_GRAY_HEX)

            set_cell_border(cell)

            set_cell_text(
                cell,
                safe_text(value),
                font_size=font_size,
                color=COLOR_DARK_GRAY
            )

    document.add_paragraph()
    return table


# =========================================================
# PREPARAÇÃO DE DATAFRAMES
# =========================================================

def select_existing_columns(df, columns):
    if df is None or df.empty:
        return df

    existing_columns = [
        column for column in columns
        if column in df.columns
    ]

    return df[existing_columns].copy()


def prepare_dataframe_for_word(df):
    if df is None or df.empty:
        return df

    prepared_df = df.copy()

    for column in prepared_df.columns:
        column_lower = str(column).lower()

        if pd.api.types.is_numeric_dtype(prepared_df[column]):

            if any(
                keyword in column_lower
                for keyword in [
                    "% cdi",
                    "cdi",
                    "taxa",
                    "rentab",
                    "alíq",
                    "aliq",
                    "alíquota",
                    "aliquota",
                    "%",
                ]
            ):
                prepared_df[column] = prepared_df[column].apply(format_percent)

            elif any(
                keyword in column_lower
                for keyword in [
                    "valor",
                    "saldo",
                    "aporte",
                    "aportado",
                    "resgate",
                    "resgatado",
                    "rendimento",
                    "ir",
                    "bruto",
                    "líquido",
                    "liquido",
                ]
            ):
                prepared_df[column] = prepared_df[column].apply(format_currency)

            else:
                prepared_df[column] = prepared_df[column].apply(format_number)

    return prepared_df


# =========================================================
# GRÁFICO DA CURVA
# =========================================================

def create_curve_chart_image(curve_df):
    if curve_df is None or curve_df.empty:
        return None

    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        chart_df = curve_df.copy()

        chart_df["Vértice"] = chart_df["Vértice"].astype(str)
        chart_df["Taxa Selic Esperada (%)"] = (
            chart_df["Taxa Selic Esperada (%)"].astype(float)
        )

        x_values = chart_df["Vértice"].tolist()
        y_values = chart_df["Taxa Selic Esperada (%)"].tolist()

        if not x_values or not y_values:
            return None

        selic_reference = y_values[0]

        fig, ax = plt.subplots(figsize=(8.2, 4.2))

        ax.plot(
            x_values,
            y_values,
            marker="o",
            linewidth=2.5,
            label="Selic esperada"
        )

        ax.axhline(
            y=selic_reference,
            linestyle="--",
            linewidth=1.5,
            label="Selic atual"
        )

        for index, value in enumerate(y_values):
            ax.annotate(
                f"{value:.2f}%",
                (x_values[index], y_values[index]),
                textcoords="offset points",
                xytext=(0, 8),
                ha="center",
                fontsize=9
            )

        ax.set_title(
            "Curva Simplificada de Juros",
            fontsize=13,
            fontweight="bold"
        )
        ax.set_xlabel("Horizonte da expectativa")
        ax.set_ylabel("Taxa Selic esperada (%)")
        ax.grid(True, axis="y", alpha=0.25)
        ax.legend(loc="best")

        image_stream = BytesIO()
        plt.tight_layout()
        fig.savefig(image_stream, format="png", dpi=180)
        plt.close(fig)

        image_stream.seek(0)
        return image_stream

    except Exception:
        return None


# =========================================================
# LEITURA DA CURVA
# =========================================================

def infer_curve_reading_from_market_intelligence(market_intelligence):
    if not market_intelligence:
        return None, None, None

    movimento_curva = market_intelligence.get("movimento_curva")
    spread_final = market_intelligence.get("spread_final")
    leitura_movimento = market_intelligence.get("leitura_movimento")

    if movimento_curva and spread_final is not None and leitura_movimento:
        return movimento_curva, spread_final, leitura_movimento

    curve_df = market_intelligence.get("curve_df")

    if curve_df is None or curve_df.empty:
        return movimento_curva, spread_final, leitura_movimento

    try:
        chart_df = curve_df.copy()

        chart_df["Taxa Selic Esperada (%)"] = (
            chart_df["Taxa Selic Esperada (%)"].astype(float)
        )

        selic_atual_referencia = chart_df["Taxa Selic Esperada (%)"].iloc[0]

        spread_final = (
            chart_df["Taxa Selic Esperada (%)"].iloc[-1]
            - selic_atual_referencia
        )

        limite_neutro = 0.10

        if spread_final > limite_neutro:
            movimento_curva = "curva abrindo"
            leitura_movimento = (
                "A curva está abrindo em relação à Selic atual. "
                "Isso indica que as expectativas de mercado apontam para juros futuros "
                "acima da taxa corrente, o que pode favorecer uma conversa consultiva "
                "sobre proteção de taxa, prazo e alternativas prefixadas, sempre conforme "
                "o perfil e a necessidade de liquidez do cliente."
            )

        elif spread_final < -limite_neutro:
            movimento_curva = "curva fechando"
            leitura_movimento = (
                "A curva está fechando em relação à Selic atual. "
                "Isso indica que as expectativas de mercado apontam para juros futuros "
                "abaixo da taxa corrente, o que reforça a importância de avaliar o risco "
                "de reinvestimento, o horizonte da aplicação e o momento adequado para "
                "travar taxas em produtos prefixados ou híbridos."
            )

        else:
            movimento_curva = "curva praticamente estável"
            leitura_movimento = (
                "A curva está praticamente estável em relação à Selic atual. "
                "Nesse cenário, a leitura consultiva deve priorizar liquidez, prazo, "
                "tributação, previsibilidade e aderência ao objetivo financeiro do cliente."
            )

        return movimento_curva, spread_final, leitura_movimento

    except Exception:
        return movimento_curva, spread_final, leitura_movimento


# =========================================================
# DRIVERS DE MERCADO QUE AFETAM O CDI
# =========================================================

def build_cdi_market_drivers_reading(market_intelligence):
    if market_intelligence is None:
        return (
            "A leitura dos drivers de mercado não foi gerada porque a inteligência "
            "de mercado não foi carregada nesta execução."
        )

    curve_shape = market_intelligence.get("curve_shape", "não informada")

    movimento_curva, spread_final, _ = (
        infer_curve_reading_from_market_intelligence(market_intelligence)
    )

    if movimento_curva is None:
        movimento_curva = "movimento não classificado"

    spread_text = "não disponível"

    if spread_final is not None:
        try:
            spread_text = (
                f"{str(round(float(spread_final), 2)).replace('.', ',')} "
                "ponto percentual"
            )
        except Exception:
            spread_text = "não disponível"

    reading = f"""
A trajetória do CDI deve ser interpretada a partir de um conjunto de fatores macroeconômicos, microeconômicos e de mercado. Nesta simulação, a curva simplificada foi classificada como {curve_shape}, com leitura de {movimento_curva} em relação à Selic atual. O spread do último vértice frente à taxa corrente é de {spread_text}.

Do ponto de vista macroeconômico, a inflação esperada é uma das variáveis centrais. Quando as expectativas de IPCA permanecem pressionadas, o Banco Central tende a ter menos espaço para reduzir a Selic, o que pode manter o CDI em patamar elevado por mais tempo. Em cenário de inflação em desaceleração, por outro lado, aumenta a probabilidade de compressão futura do CDI e cresce a importância de avaliar risco de reinvestimento.

A política monetária também deve ser lida em conjunto com câmbio, atividade econômica e risco fiscal. A depreciação cambial pode pressionar preços e limitar cortes de juros. A atividade econômica aquecida pode sustentar inflação de serviços, enquanto desaceleração do PIB pode ampliar o espaço para redução da Selic. A piora da percepção fiscal pode abrir a curva de juros, aumentando o prêmio exigido para prazos mais longos, mesmo sem alteração imediata da taxa básica.

No mercado de crédito privado, movimentos de liquidação de títulos, resgates em fundos ou abertura de spreads podem alterar a relação entre retorno esperado, risco e liquidez. Em momentos de estresse, os ativos privados podem exigir maior prêmio, mas também demandam análise mais cuidadosa de emissor, prazo, garantias, liquidez secundária e risco de marcação.

Cenários de juros elevados tendem a atrair investidores para posições de renda fixa, especialmente produtos pós-fixados, Tesouro Selic, CDBs, LCIs, LCAs e fundos DI. Esse fluxo pode reforçar a demanda por instrumentos atrelados ao CDI, mas também pode reduzir prêmios em ativos mais disputados. Quando há expectativa consistente de queda da Selic, produtos prefixados, híbridos ou indexados à inflação podem ganhar relevância na conversa consultiva, desde que o horizonte de investimento e a tolerância a risco sejam compatíveis.

Assim, a análise consultiva não deve considerar apenas a taxa projetada. Ela deve ponderar liquidez, prazo, tributação, risco de crédito, qualidade do emissor, comportamento da inflação, risco fiscal, curva de juros, possibilidade de reinvestimento a taxas menores e risco de oportunidade caso o cenário de juros se altere.
"""

    return reading.strip()


# =========================================================
# FUNÇÃO PRINCIPAL
# =========================================================

def generate_word_report(
    client_name,
    advisor_name,
    simulation_mode,
    start_date,
    end_date,
    months,
    initial_amount,
    annual_cdi_rate,
    selic_rate,
    tr_rate,
    cdb_percentage,
    lci_lca_percentage,
    treasury_percentage,
    treasury_annual_fee,
    fund_percentage,
    fund_annual_fee,
    fund_type=None,
    apply_come_cotas=True,
    fund_come_cotas_tax=0.0,
    fund_redemption_tax=0.0,
    fund_total_tax=0.0,
    fund_admin_fee_impact=0.0,
    fund_net_final_amount=0.0,
    fund_net_return_percentage=0.0,
    fund_come_cotas_rate=0.0,
    fund_final_ir_rate=0.0,
    comparison_df=None,
    cashflow_df=None,
    monthly_df=None,
    consultive_analysis=None,
    market_intelligence=None,
    include_dividend_scenario=False,
    monthly_dividends=0.0,
    same_payer_dividends=True,
    months_with_dividends=0,
    taxable_monthly_dividends=0.0,
    estimated_monthly_dividend_ir=0.0,
    integrated_tax_scenario=None,
    estimated_annual_dividend_ir=0.0,
    report_options=None,
    panorama_mercado=None,
):

    if report_options is None:
        report_options = {
            "visao_geral": True,
            "premissas": True,
            "comparativo": True,
            "calendario": True,
            "resumo_mensal": True,
            "leitura_consultiva": True,
            "inteligencia_mercado": True,
            "curva_juros": True,
            "leitura_foresight": True,
            "aviso_tecnico": True,
        }

    document = Document()

    styles = document.styles

    styles["Normal"].font.name = FONT_BODY
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = COLOR_DARK_GRAY

    styles["Heading 1"].font.name = FONT_TITLE
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = COLOR_NAVY

    styles["Heading 2"].font.name = FONT_BODY
    styles["Heading 2"].font.size = Pt(11)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = COLOR_BLUE

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


    # =========================================================
    # FAIXA SUPERIOR
    # =========================================================

    top_band = document.add_table(rows=1, cols=1)
    top_band.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_band.style = "Table Grid"

    top_cell = top_band.rows[0].cells[0]
    shade_cell(top_cell, COLOR_NAVY_HEX)
    set_cell_border(top_cell, color=COLOR_NAVY_HEX, size="0")

    top_cell.text = ""
    top_paragraph = top_cell.paragraphs[0]
    top_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    top_run = top_paragraph.add_run("SIMULADOR ESTRATÉGICO DE INVESTIMENTOS")
    top_run.bold = True
    top_run.font.name = FONT_BODY
    top_run.font.size = Pt(8)
    top_run.font.color.rgb = COLOR_WHITE

    document.add_paragraph()

    # =========================================================
    # CABEÇALHO EXECUTIVO
    # =========================================================

    header_table = document.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.style = "Table Grid"

    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]

    shade_cell(left_cell, "FFFFFF")
    shade_cell(right_cell, "FFFFFF")
    set_cell_border(left_cell)
    set_cell_border(right_cell)

    left_cell.text = ""
    right_cell.text = ""

    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title_run = left_paragraph.add_run(
        "RELATÓRIO CONSULTIVO DE INVESTIMENTOS"
    )
    title_run.bold = True
    title_run.font.name = FONT_TITLE
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = COLOR_NAVY

    subtitle_paragraph = left_cell.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run(
        "Simulação, análise comparativa e leitura estratégica de mercado"
    )
    subtitle_run.font.name = FONT_BODY
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = COLOR_BLUE

    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    badge_run = right_paragraph.add_run("EXTRATO SIMULADO")
    badge_run.bold = True
    badge_run.font.name = FONT_BODY
    badge_run.font.size = Pt(9)
    badge_run.font.color.rgb = COLOR_BLUE

    document.add_paragraph()

    metadata_rows = [
        ("Cliente", client_name),
        ("Banker / Assessor", advisor_name),
        ("Período", f"{start_date} a {end_date}"),
        ("Modo de simulação", simulation_mode),
    ]

    add_key_value_table(document, metadata_rows)

    section_number = 1

    # =========================================================
    # CÁLCULO DA MELHOR ALTERNATIVA
    # =========================================================

    best_product = ""
    best_net_value = 0
    best_net_profit = 0
    best_net_return = 0
    worst_net_value = 0
    difference_best_worst = 0

    if comparison_df is not None and not comparison_df.empty:
        try:
            value_column = "Valor Líquido"
            profit_column = "Rendimento Líquido"
            return_column = "Rentab. Líq. Período (%)"
            product_column = "Produto"

            sorted_df = comparison_df.sort_values(
                by=value_column,
                ascending=False
            )

            best_row = sorted_df.iloc[0]
            worst_row = sorted_df.iloc[-1]

            best_product = safe_text(best_row.get(product_column, ""))
            best_net_value = float(best_row.get(value_column, 0))
            best_net_profit = float(best_row.get(profit_column, 0))
            best_net_return = float(best_row.get(return_column, 0))
            worst_net_value = float(worst_row.get(value_column, 0))
            difference_best_worst = best_net_value - worst_net_value

        except Exception:
            pass

    # =========================================================
    # PAINEL EXECUTIVO
    # =========================================================

    if report_options.get("visao_geral", True):
        add_section_heading(
            document,
            section_number,
            "Painel Executivo"
        )
        section_number += 1

        movimento_curva_card = "Não informado"

        if market_intelligence is not None:
            movimento_curva_tmp, _, _ = (
                infer_curve_reading_from_market_intelligence(
                    market_intelligence
                )
            )

            if movimento_curva_tmp:
                movimento_curva_card = movimento_curva_tmp.title()

        executive_cards = [
            ("Valor líquido final", format_currency(best_net_value)),
            ("Melhor alternativa", best_product),
            ("Rentabilidade líquida", format_percent(best_net_return)),
            ("Ganho líquido projetado", format_currency(best_net_profit)),
            ("Movimento da curva", movimento_curva_card),
        ]

        cards_table = document.add_table(rows=1, cols=5)
        cards_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cards_table.style = "Table Grid"

        for index, (label, value) in enumerate(executive_cards):
            cell = cards_table.rows[0].cells[index]

            shade_cell(cell, COLOR_LIGHT_BLUE_HEX)
            set_cell_border(cell, color=COLOR_CARD_BORDER_HEX, size="6")

            cell.text = ""

            p_label = cell.paragraphs[0]
            p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_label = p_label.add_run(label.upper())
            run_label.bold = True
            run_label.font.name = FONT_BODY
            run_label.font.size = Pt(7)
            run_label.font.color.rgb = COLOR_BLUE

            p_value = cell.add_paragraph()
            p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_value = p_value.add_run(value)
            run_value.bold = True
            run_value.font.name = FONT_TITLE
            run_value.font.size = Pt(11)
            run_value.font.color.rgb = COLOR_NAVY

        document.add_paragraph()

    # =========================================================
    # PREMISSAS UTILIZADAS
    # =========================================================

    if report_options.get("premissas", True):
        add_section_heading(
            document,
            section_number,
            "Premissas Utilizadas"
        )
        section_number += 1

        premises_rows = [
            ("Valor inicial", format_currency(initial_amount)),
            ("CDI anual estimado", format_percent(annual_cdi_rate)),
            ("Selic meta estimada", format_percent(selic_rate)),
            ("TR anual estimada", format_percent(tr_rate)),
            ("CDB / LC", f"{format_percent(cdb_percentage)} do CDI"),
            ("LCI / LCA", f"{format_percent(lci_lca_percentage)} do CDI"),
            ("Tesouro Selic", f"{format_percent(treasury_percentage)} do CDI"),
            ("Custo anual Tesouro Selic", format_percent(treasury_annual_fee)),
            ("Fundo DI", f"{format_percent(fund_percentage)} do CDI"),
            ("Taxa de administração Fundo DI", format_percent(fund_annual_fee)),
        ]

        if cashflow_df is not None and not cashflow_df.empty:
            try:
                total_aportado = cashflow_df[
                    cashflow_df["Tipo"].str.lower() == "aporte"
                ]["Valor"].astype(float).sum()

                total_resgatado = cashflow_df[
                    cashflow_df["Tipo"].str.lower() == "resgate"
                ]["Valor"].astype(float).sum()

                premises_rows.append(
                    ("Total aportado", format_currency(total_aportado))
                )

                premises_rows.append(
                    ("Total resgatado", format_currency(total_resgatado))
                )

            except Exception:
                pass

        add_key_value_table(document, premises_rows)

    # =========================================================
    # COMPARATIVO DOS PRODUTOS
    # =========================================================

    if report_options.get("comparativo", True):
        add_section_heading(
            document,
            section_number,
            "Comparativo dos Produtos"
        )
        section_number += 1

        add_paragraph(
            document,
            "As tabelas a seguir organizam a comparação entre os produtos simulados "
            "em duas leituras complementares: resultado bruto e resultado líquido. "
            "Essa separação melhora a leitura do impacto da tributação, dos custos "
            "e da rentabilidade final projetada."
        )

        gross_columns = [
            "Produto",
            "% CDI",
            "Taxa Efetiva a.a. (%)",
            "Valor Inicial",
            "Total Aportado",
            "Total Resgatado",
            "Valor Bruto",
            "Rendimento Bruto",
            "Rentab. Bruta Período (%)",
            "Rentab. Bruta Mês (%)",
            "Rentab. Bruta Ano (%)",
        ]

        gross_df = select_existing_columns(
            comparison_df,
            gross_columns
        )

        net_columns = [
            "Produto",
            "Valor Bruto",
            "Rendimento Bruto",
            "IR",
            "Alíq. IR (%)",
            "Valor Líquido",
            "Rendimento Líquido",
            "Rentab. Líq. Período (%)",
            "Rentab. Líq. Mês (%)",
            "Rentab. Líq. Ano (%)",
            "Tributável",
        ]

        net_df = select_existing_columns(
            comparison_df,
            net_columns
        )

        if gross_df is not None and not gross_df.empty:
            add_subheading(
                document,
                "Comparativo de Rentabilidade Bruta"
            )

            gross_word_df = prepare_dataframe_for_word(gross_df)

            add_dataframe_table(
                document,
                gross_word_df,
                font_size=7
            )

        if net_df is not None and not net_df.empty:
            add_subheading(
                document,
                "Comparativo de Rentabilidade Líquida"
            )

            net_word_df = prepare_dataframe_for_word(net_df)

            net_word_df = net_word_df.rename(
                columns={
                    "IR": "IR (R$)",
                    "Alíq. IR (%)": "Alíquota IR",
                    "Rendimento Bruto": "Rend. Bruto",
                    "Rendimento Líquido": "Rend. Líquido",
                    "Rentab. Líq. Período (%)": "Rentab. Período",
                    "Rentab. Líq. Mês (%)": "Rentab. Mês",
                    "Rentab. Líq. Ano (%)": "Rentab. Ano",
                }
            )

            add_dataframe_table(
                document,
                net_word_df,
                font_size=6
            )

            # =========================================================
            # CENÁRIO ANUAL: DIVIDENDOS E TRIBUTAÇÃO MÍNIMA
            # =========================================================

            if integrated_tax_scenario and integrated_tax_scenario.get(
                "include_dividend_scenario"
            ):
                add_subheading(
                    document,
                    "Cenário Anual: Dividendos, Tributação Mínima e Alocação"
                )

                monthly_dividends = integrated_tax_scenario.get(
                    "monthly_dividends",
                    0.0
                )
                months_with_dividends = integrated_tax_scenario.get(
                    "months_with_dividends",
                    0
                )
                taxable_monthly_dividends = integrated_tax_scenario.get(
                    "taxable_monthly_dividends",
                    0.0
                )
                estimated_monthly_dividend_ir = integrated_tax_scenario.get(
                    "estimated_monthly_dividend_ir",
                    0.0
                )
                estimated_annual_dividend_ir = integrated_tax_scenario.get(
                    "estimated_annual_dividend_ir",
                    0.0
                )
                annual_total_income = integrated_tax_scenario.get(
                    "annual_total_income",
                    0.0
                )
                minimum_tax_rate = integrated_tax_scenario.get(
                    "minimum_tax_rate",
                    0.0
                )
                minimum_tax_due = integrated_tax_scenario.get(
                    "minimum_tax_due",
                    0.0
                )
                isolated_winner_product = integrated_tax_scenario.get(
                    "isolated_winner_product",
                    ""
                )
                isolated_winner_net_value = integrated_tax_scenario.get(
                    "isolated_winner_net_value",
                    0.0
                )
                adjusted_winner_product = integrated_tax_scenario.get(
                    "adjusted_winner_product",
                    ""
                )
                adjusted_winner_value = integrated_tax_scenario.get(
                    "adjusted_winner_value",
                    0.0
                )
                net_difference_lci_vs_cdb = integrated_tax_scenario.get(
                    "net_difference_lci_vs_cdb",
                    0.0
                )
                fiscal_effect_cdb_vs_lci = integrated_tax_scenario.get(
                    "fiscal_effect_cdb_vs_lci",
                    0.0
                )

                document.add_paragraph(
                    "Esta seção apresenta uma leitura complementar da decisão "
                    "de alocação quando o cliente possui recebimento relevante "
                    "de lucros e dividendos. A análise não considera compensação "
                    "automática mensal entre produtos financeiros e dividendos. "
                    "O objetivo é comparar o resultado líquido isolado dos "
                    "produtos com uma visão anual da tributação mínima estimada."
                )

                document.add_paragraph(
                    f"Nesta simulação, o cliente informou dividendos mensais "
                    f"estimados de R$ {monthly_dividends:,.2f}, durante "
                    f"{months_with_dividends} mês(es) no ano. A base mensal "
                    f"considerada para retenção foi de "
                    f"R$ {taxable_monthly_dividends:,.2f}. O IRRF mensal "
                    f"estimado sobre dividendos foi de "
                    f"R$ {estimated_monthly_dividend_ir:,.2f}, resultando em "
                    f"IRRF anual estimado de "
                    f"R$ {estimated_annual_dividend_ir:,.2f}."
                )

                document.add_paragraph(
                    f"A renda anual total estimada informada foi de "
                    f"R$ {annual_total_income:,.2f}. Com base nesse valor, "
                    f"a alíquota mínima estimada foi de "
                    f"{minimum_tax_rate * 100:.2f}%, gerando IR mínimo anual "
                    f"estimado de R$ {minimum_tax_due:,.2f}."
                )

                document.add_paragraph(
                    f"No resultado líquido isolado da aplicação, o produto "
                    f"com melhor desempenho foi {isolated_winner_product}, "
                    f"com valor líquido estimado de "
                    f"R$ {isolated_winner_net_value:,.2f}."
                )

                document.add_paragraph(
                    f"Na leitura fiscal anual, o produto com melhor valor "
                    f"comparável no cenário foi {adjusted_winner_product}, "
                    f"com valor estimado de R$ {adjusted_winner_value:,.2f}. "
                    f"Essa métrica considera o efeito fiscal potencial dos "
                    f"produtos tributados frente à LCI/LCA, sem tratar esse "
                    f"valor como saldo financeiro automático."
                )

                document.add_paragraph(
                    f"A diferença líquida da LCI/LCA em relação ao CDB/LC "
                    f"foi de R$ {net_difference_lci_vs_cdb:,.2f}. O efeito "
                    f"fiscal potencial do CDB/LC frente à LCI/LCA foi estimado "
                    f"em R$ {fiscal_effect_cdb_vs_lci:,.2f}. Quando esse "
                    f"efeito fiscal potencial supera a diferença líquida entre "
                    f"os produtos, o CDB/LC merece análise consultiva mais "
                    f"aprofundada no cenário anual."
                )

                document.add_paragraph(
                    "A leitura final deve separar dois planos de decisão. No "
                    "plano da aplicação isolada, prevalece o produto com maior "
                    "valor líquido projetado. No plano fiscal anual, produtos "
                    "tributados podem ganhar relevância quando o imposto retido "
                    "no produto reduz eventual saldo adicional de tributação "
                    "mínima. A aplicação prática depende da situação fiscal "
                    "completa do cliente e deve ser validada por contador ou "
                    "especialista tributário."
                )

        add_subheading(
            document,
            "Detalhamento Tributário do Fundo DI"
        )

        fund_tax_text = (
            f"O Fundo DI foi simulado como {fund_type}. "
            f"A aplicação do come-cotas foi considerada como "
            f"{'ativa' if apply_come_cotas else 'inativa'} na simulação. "
            f"A alíquota de come-cotas utilizada foi de "
            f"{fund_come_cotas_rate * 100:.1f}% e a alíquota final de IR "
            f"estimada para o resgate foi de {fund_final_ir_rate * 100:.1f}%."
        )

        add_paragraph(
            document,
            fund_tax_text
        )

        add_paragraph(
            document,
            f"No cenário simulado, o impacto estimado da taxa de administração "
            f"foi de {format_currency(fund_admin_fee_impact)}. "
            f"O come-cotas estimado foi de {format_currency(fund_come_cotas_tax)}, "
            f"enquanto o IR complementar no resgate foi de "
            f"{format_currency(fund_redemption_tax)}. O IR total estimado "
            f"para o Fundo DI foi de {format_currency(fund_total_tax)}."
        )

        add_paragraph(
            document,
            f"O valor líquido final estimado do Fundo DI foi de "
            f"{format_currency(fund_net_final_amount)}, com rentabilidade líquida "
            f"de {format_percent(fund_net_return_percentage)} no período."
        )

        add_paragraph(
            document,
            "Essa leitura é importante porque fundos sujeitos ao come-cotas não "
            "devem ser comparados apenas pelo percentual do CDI. A análise deve "
            "considerar a classificação fiscal do fundo, a taxa de administração, "
            "a antecipação semestral do Imposto de Renda, o prazo de permanência "
            "e eventual complemento tributário no resgate."
        )
            

    # =========================================================
    # CENÁRIO ANUAL: DIVIDENDOS, TRIBUTAÇÃO MÍNIMA E ALOCAÇÃO
    # =========================================================

    if integrated_tax_scenario and integrated_tax_scenario.get(
        "include_dividend_scenario"
    ):
        add_section_heading(
            document,
            section_number,
            "Cenário Anual: Dividendos, Tributação Mínima e Alocação"
        )
        section_number += 1

        monthly_dividends = integrated_tax_scenario.get(
            "monthly_dividends",
            0.0
        )
        months_with_dividends = integrated_tax_scenario.get(
            "months_with_dividends",
            0
        )
        taxable_monthly_dividends = integrated_tax_scenario.get(
            "taxable_monthly_dividends",
            0.0
        )
        estimated_monthly_dividend_ir = integrated_tax_scenario.get(
            "estimated_monthly_dividend_ir",
            0.0
        )
        estimated_annual_dividend_ir = integrated_tax_scenario.get(
            "estimated_annual_dividend_ir",
            0.0
        )
        annual_total_income = integrated_tax_scenario.get(
            "annual_total_income",
            0.0
        )
        minimum_tax_rate = integrated_tax_scenario.get(
            "minimum_tax_rate",
            0.0
        )
        minimum_tax_due = integrated_tax_scenario.get(
            "minimum_tax_due",
            0.0
        )
        isolated_winner_product = integrated_tax_scenario.get(
            "isolated_winner_product",
            ""
        )
        isolated_winner_net_value = integrated_tax_scenario.get(
            "isolated_winner_net_value",
            0.0
        )
        adjusted_winner_product = integrated_tax_scenario.get(
            "adjusted_winner_product",
            ""
        )
        adjusted_winner_value = integrated_tax_scenario.get(
            "adjusted_winner_value",
            0.0
        )
        net_difference_lci_vs_cdb = integrated_tax_scenario.get(
            "net_difference_lci_vs_cdb",
            0.0
        )
        fiscal_effect_cdb_vs_lci = integrated_tax_scenario.get(
            "fiscal_effect_cdb_vs_lci",
            0.0
        )
        annual_tax_scenario_records = integrated_tax_scenario.get(
            "annual_tax_scenario_records",
            []
        )

        add_paragraph(
            document,
            "Esta seção compara a decisão de alocação em dois planos: o "
            "resultado líquido isolado dos produtos e a leitura fiscal anual "
            "associada à tributação mínima. A análise não considera compensação "
            "automática mensal entre produtos financeiros e dividendos. O objetivo "
            "é avaliar se impostos já pagos ou retidos podem reduzir eventual "
            "saldo adicional de tributação mínima, quando aplicável ao caso concreto."
        )

        add_paragraph(
            document,
            f"Nesta simulação, o cliente informou dividendos mensais estimados "
            f"de {format_currency(monthly_dividends)}, durante "
            f"{months_with_dividends} mês(es) no ano. A base mensal considerada "
            f"para retenção foi de {format_currency(taxable_monthly_dividends)}. "
            f"O IRRF mensal estimado sobre dividendos foi de "
            f"{format_currency(estimated_monthly_dividend_ir)}, resultando em "
            f"IRRF anual estimado de "
            f"{format_currency(estimated_annual_dividend_ir)}."
        )

        add_paragraph(
            document,
            f"A renda anual total estimada informada foi de "
            f"{format_currency(annual_total_income)}. Com base nesse valor, "
            f"a alíquota mínima estimada foi de "
            f"{format_percent(minimum_tax_rate * 100)}, gerando IR mínimo anual "
            f"estimado de {format_currency(minimum_tax_due)}."
        )

        if annual_tax_scenario_records:
            annual_tax_word_df = pd.DataFrame(annual_tax_scenario_records)

            columns_to_show = [
                "Produto",
                "Valor líquido da aplicação",
                "Rentab. líquida isolada (%)",
                "IR do produto",
                "IRRF dividendos",
                "Saldo adicional estimado",
                "Efeito fiscal potencial vs LCI/LCA",
                "Valor comparável no cenário",
                "Rentab. comparável no cenário (%)",
            ]

            annual_tax_word_df = select_existing_columns(
                annual_tax_word_df,
                columns_to_show
            )

            annual_tax_word_df = prepare_dataframe_for_word(
                annual_tax_word_df
            )

            add_dataframe_table(
                document,
                annual_tax_word_df,
                font_size=6
            )

        add_paragraph(
            document,
            f"No resultado líquido isolado da aplicação, o produto com melhor "
            f"desempenho foi {isolated_winner_product}, com valor líquido "
            f"estimado de {format_currency(isolated_winner_net_value)}."
        )

        add_paragraph(
            document,
            f"Na leitura fiscal anual, o produto com maior valor comparável "
            f"no cenário foi {adjusted_winner_product}, com valor estimado de "
            f"{format_currency(adjusted_winner_value)}. Essa métrica não representa "
            f"saldo financeiro automático. Ela compara o valor líquido da aplicação "
            f"com o efeito fiscal potencial dos produtos tributados em relação "
            f"à LCI/LCA."
        )

        add_paragraph(
            document,
            f"A diferença líquida da LCI/LCA em relação ao CDB/LC foi de "
            f"{format_currency(net_difference_lci_vs_cdb)}. O efeito fiscal "
            f"potencial do CDB/LC frente à LCI/LCA foi estimado em "
            f"{format_currency(fiscal_effect_cdb_vs_lci)}. Quando esse efeito "
            f"fiscal potencial supera a diferença líquida entre os produtos, "
            f"o CDB/LC merece análise consultiva mais aprofundada no cenário anual."
        )

        add_paragraph(
            document,
            "A decisão final deve separar a eficiência da aplicação isolada "
            "da eficiência fiscal anual. LCI e LCA podem continuar superiores "
            "quando a prioridade é rentabilidade líquida isenta. Produtos "
            "tributados, como CDB/LC, Tesouro Selic e Fundo DI, podem ganhar "
            "relevância quando o imposto retido no produto reduz eventual saldo "
            "adicional de tributação mínima. A aplicação prática depende da "
            "situação fiscal completa do cliente e deve ser validada por "
            "contador ou especialista tributário."
        )

        
    # =========================================================
    # INTELIGÊNCIA DE MERCADO E FORESIGHT
    # =========================================================

    if report_options.get("inteligencia_mercado", True):
        add_section_heading(
            document,
            section_number,
            "Inteligência de Mercado e Foresight"
        )
        section_number += 1

        if market_intelligence is None:
            add_paragraph(
                document,
                "A inteligência de mercado não foi carregada nesta execução. "
                "Para incluir dados Bacen/Focus/Foresight, ative o módulo no app, "
                "clique em 'Atualizar inteligência de mercado' e gere o relatório novamente."
            )
        else:
            add_paragraph(
                document,
                "Os dados públicos do Banco Central e as expectativas do Boletim Focus "
                "foram incorporados à simulação para apoiar a leitura macroeconômica "
                "e qualificar a análise consultiva. Esta camada não altera os cálculos "
                "da simulação CDI, mas amplia a contextualização da decisão."
            )

            curve_shape = market_intelligence.get("curve_shape")
            market_reading = market_intelligence.get("reading")
            llm_foresight = (
                market_intelligence.get("llm_foresight")
                or market_intelligence.get("llm_reading")
            )

            if curve_shape:
                add_paragraph(
                    document,
                    f"Classificação da curva: {curve_shape}.",
                    bold=True
                )

            if market_reading:
                add_paragraph(
                    document,
                    market_reading
                )

            add_subheading(
                document,
                "Drivers de Mercado que Afetam o CDI"
            )

            cdi_market_drivers_reading = build_cdi_market_drivers_reading(
                market_intelligence
            )

            add_paragraph(
                document,
                cdi_market_drivers_reading
            )

            if llm_foresight:
                add_subheading(
                    document,
                    "Leitura Foresight Assistida por LLM"
                )

                add_paragraph(
                    document,
                    llm_foresight
                )

    # =========================================================
    # CURVA SIMPLIFICADA DE JUROS
    # =========================================================

    if report_options.get("curva_juros", True):
        add_section_heading(
            document,
            section_number,
            "Curva Simplificada de Juros"
        )
        section_number += 1

        if market_intelligence is None:
            add_paragraph(
                document,
                "A curva simplificada de juros não foi carregada nesta execução. "
                "Para gerar esta seção com gráfico e tabela, atualize a inteligência "
                "de mercado antes de baixar o relatório."
            )
        else:
            curve_df = market_intelligence.get("curve_df")

            if curve_df is None or curve_df.empty:
                add_paragraph(
                    document,
                    "A curva simplificada de juros não está disponível nesta execução."
                )
            else:
                chart_image = create_curve_chart_image(curve_df)

                if chart_image is not None:
                    paragraph = document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(chart_image, width=Inches(6.4))
                    document.add_paragraph()
                else:
                    add_paragraph(
                        document,
                        "O gráfico da curva não pôde ser gerado nesta execução. "
                        "A tabela técnica da curva segue apresentada abaixo."
                    )

                add_subheading(document, "Tabela técnica da curva")

                curve_word_df = prepare_dataframe_for_word(curve_df)

                add_dataframe_table(
                    document,
                    curve_word_df,
                    font_size=8
                )

    # =========================================================
    # LEITURA FORESIGHT DA CURVA
    # =========================================================

    if report_options.get("leitura_foresight", True):
        add_section_heading(
            document,
            section_number,
            "Leitura Foresight da Curva"
        )
        section_number += 1

        if market_intelligence is None:
            add_paragraph(
                document,
                "A leitura foresight da curva não foi gerada porque a inteligência "
                "de mercado não foi carregada nesta execução."
            )
        else:
            movimento_curva, spread_final, leitura_movimento = (
                infer_curve_reading_from_market_intelligence(market_intelligence)
            )

            if movimento_curva:
                add_paragraph(
                    document,
                    f"Em relação à Selic atual, a estrutura observada indica "
                    f"{movimento_curva}.",
                    bold=True
                )

                if spread_final is not None:
                    add_paragraph(
                        document,
                        f"O spread do último vértice frente à taxa corrente é de "
                        f"{str(round(float(spread_final), 2)).replace('.', ',')} "
                        f"ponto percentual."
                    )

                if leitura_movimento:
                    add_paragraph(
                        document,
                        leitura_movimento
                    )
            else:
                add_paragraph(
                    document,
                    "Não foi possível classificar a leitura foresight da curva nesta execução."
                )

    # =========================================================
    # CALENDÁRIO DE MOVIMENTAÇÕES
    # =========================================================

    if report_options.get("calendario", True):
        if cashflow_df is not None and not cashflow_df.empty:
            add_section_heading(
                document,
                section_number,
                "Calendário de Movimentações"
            )
            section_number += 1

            cashflow_word_df = prepare_dataframe_for_word(cashflow_df)

            add_dataframe_table(
                document,
                cashflow_word_df,
                font_size=8
            )
        else:
            add_section_heading(
                document,
                section_number,
                "Calendário de Movimentações"
            )
            section_number += 1

            add_paragraph(
                document,
                "Não há movimentações de aporte ou resgate informadas para esta simulação."
            )

    # =========================================================
    # RESUMO MENSAL
    # =========================================================

    if report_options.get("resumo_mensal", True):
        add_section_heading(
            document,
            section_number,
            "Apêndice Técnico — Resumo Mensal"
        )
        section_number += 1

        if monthly_df is not None and not monthly_df.empty:
            monthly_word_df = prepare_dataframe_for_word(monthly_df)

            add_dataframe_table(
                document,
                monthly_word_df,
                font_size=6
            )
        else:
            add_paragraph(
                document,
                "O resumo mensal não está disponível para esta simulação."
            )

    # =========================================================
    # LEITURA CONSULTIVA
    # =========================================================

    if report_options.get("leitura_consultiva", True):
        add_section_heading(
            document,
            section_number,
            "Leitura Consultiva"
        )
        section_number += 1

        if consultive_analysis:
            add_paragraph(
                document,
                consultive_analysis
            )
        else:
            add_paragraph(
                document,
                f"A simulação indica que, para as premissas informadas, "
                f"a alternativa com maior valor líquido projetado é {best_product}."
            )

            add_paragraph(
                document,
                f"O valor líquido estimado para essa alternativa é de "
                f"{format_currency(best_net_value)}, com rendimento líquido aproximado "
                f"de {format_currency(best_net_profit)} no período. A rentabilidade "
                f"líquida projetada no período é de {format_percent(best_net_return)}."
            )

            if difference_best_worst:
                add_paragraph(
                    document,
                    f"A diferença entre a melhor alternativa e a alternativa com menor "
                    f"valor líquido, neste cenário, é de aproximadamente "
                    f"{format_currency(difference_best_worst)}."
                )

            add_paragraph(
                document,
                "A comparação é especialmente relevante porque produtos com percentuais "
                "diferentes do CDI podem ter resultados líquidos próximos ou até superiores, "
                "dependendo da tributação, da isenção fiscal, das taxas e do prazo da aplicação."
            )

            add_paragraph(
                document,
                "Esta simulação deve ser utilizada como apoio à análise consultiva, "
                "considerando perfil do cliente, liquidez, objetivo financeiro, horizonte "
                "de investimento e adequação do produto."
            )

    # =========================================================
    # AVISO TÉCNICO
    # =========================================================

    if report_options.get("aviso_tecnico", True):
        add_section_heading(
            document,
            section_number,
            "Aviso Técnico"
        )
        section_number += 1

        add_paragraph(
            document,
            "Os valores apresentados são estimativas para fins de simulação. "
            "Este relatório não constitui promessa de rentabilidade, oferta, "
            "recomendação individualizada ou garantia de resultado. A análise final "
            "deve considerar perfil do cliente, adequação, liquidez, risco, tributação, "
            "condições de mercado e normas aplicáveis."
        )

    # =========================================================
    # RODAPÉ
    # =========================================================

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_run = footer.add_run(
        "Relatório gerado automaticamente pelo Simulador Estratégico de Investimentos | "
        "Uso interno e consultivo"
    )
    footer_run.font.name = FONT_BODY
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = COLOR_MUTED_GRAY


    # =========================================================
    # CENÁRIO ANUAL: DIVIDENDOS, TRIBUTAÇÃO MÍNIMA E ALOCAÇÃO
    # =========================================================

    if integrated_tax_scenario and integrated_tax_scenario.get(
        "include_dividend_scenario"
    ):
        add_subheading(
            document,
            "Cenário Anual: Dividendos, Tributação Mínima e Alocação"
        )

        monthly_dividends = integrated_tax_scenario.get(
            "monthly_dividends",
            0.0
        )
        months_with_dividends = integrated_tax_scenario.get(
            "months_with_dividends",
            0
        )
        taxable_monthly_dividends = integrated_tax_scenario.get(
            "taxable_monthly_dividends",
            0.0
        )
        estimated_monthly_dividend_ir = integrated_tax_scenario.get(
            "estimated_monthly_dividend_ir",
            0.0
        )
        estimated_annual_dividend_ir = integrated_tax_scenario.get(
            "estimated_annual_dividend_ir",
            0.0
        )
        annual_total_income = integrated_tax_scenario.get(
            "annual_total_income",
            0.0
        )
        minimum_tax_rate = integrated_tax_scenario.get(
            "minimum_tax_rate",
            0.0
        )
        minimum_tax_due = integrated_tax_scenario.get(
            "minimum_tax_due",
            0.0
        )
        isolated_winner_product = integrated_tax_scenario.get(
            "isolated_winner_product",
            ""
        )
        isolated_winner_net_value = integrated_tax_scenario.get(
            "isolated_winner_net_value",
            0.0
        )
        adjusted_winner_product = integrated_tax_scenario.get(
            "adjusted_winner_product",
            ""
        )
        adjusted_winner_value = integrated_tax_scenario.get(
            "adjusted_winner_value",
            0.0
        )
        net_difference_lci_vs_cdb = integrated_tax_scenario.get(
            "net_difference_lci_vs_cdb",
            0.0
        )
        fiscal_effect_cdb_vs_lci = integrated_tax_scenario.get(
            "fiscal_effect_cdb_vs_lci",
            0.0
        )
        annual_tax_scenario_records = integrated_tax_scenario.get(
            "annual_tax_scenario_records",
            []
        )

        document.add_paragraph(
            "Esta seção compara a decisão de alocação em dois planos: o "
            "resultado líquido isolado dos produtos e a leitura fiscal anual "
            "associada à tributação mínima. A análise não considera compensação "
            "automática mensal entre produtos financeiros e dividendos. O objetivo "
            "é avaliar se impostos já pagos ou retidos podem reduzir eventual "
            "saldo adicional de tributação mínima, quando aplicável ao caso concreto."
        )

        document.add_paragraph(
            f"Nesta simulação, o cliente informou dividendos mensais estimados "
            f"de R$ {monthly_dividends:,.2f}, durante {months_with_dividends} "
            f"mês(es) no ano. A base mensal considerada para retenção foi de "
            f"R$ {taxable_monthly_dividends:,.2f}. O IRRF mensal estimado sobre "
            f"dividendos foi de R$ {estimated_monthly_dividend_ir:,.2f}, "
            f"resultando em IRRF anual estimado de "
            f"R$ {estimated_annual_dividend_ir:,.2f}."
        )

        document.add_paragraph(
            f"A renda anual total estimada informada foi de "
            f"R$ {annual_total_income:,.2f}. Com base nesse valor, a alíquota "
            f"mínima estimada foi de {minimum_tax_rate * 100:.2f}%, gerando "
            f"IR mínimo anual estimado de R$ {minimum_tax_due:,.2f}."
        )

        if annual_tax_scenario_records:
            annual_tax_word_df = pd.DataFrame(annual_tax_scenario_records)

            columns_to_show = [
                "Produto",
                "Valor líquido da aplicação",
                "Rentab. líquida isolada (%)",
                "IR do produto",
                "IRRF dividendos",
                "Saldo adicional estimado",
                "Efeito fiscal potencial vs LCI/LCA",
                "Valor comparável no cenário",
                "Rentab. comparável no cenário (%)",
            ]

            annual_tax_word_df = select_existing_columns(
                annual_tax_word_df,
                columns_to_show
            )

            annual_tax_word_df = prepare_dataframe_for_word(
                annual_tax_word_df
            )

            add_dataframe_table(
                document,
                annual_tax_word_df,
                font_size=6
            )

        document.add_paragraph(
            f"No resultado líquido isolado da aplicação, o produto com melhor "
            f"desempenho foi {isolated_winner_product}, com valor líquido "
            f"estimado de R$ {isolated_winner_net_value:,.2f}."
        )

        document.add_paragraph(
            f"Na leitura fiscal anual, o produto com maior valor comparável "
            f"no cenário foi {adjusted_winner_product}, com valor estimado de "
            f"R$ {adjusted_winner_value:,.2f}. Essa métrica não representa saldo "
            f"financeiro automático. Ela compara o valor líquido da aplicação "
            f"com o efeito fiscal potencial dos produtos tributados em relação "
            f"à LCI/LCA."
        )

        document.add_paragraph(
            f"A diferença líquida da LCI/LCA em relação ao CDB/LC foi de "
            f"R$ {net_difference_lci_vs_cdb:,.2f}. O efeito fiscal potencial "
            f"do CDB/LC frente à LCI/LCA foi estimado em "
            f"R$ {fiscal_effect_cdb_vs_lci:,.2f}. Como esse efeito fiscal "
            f"potencial supera a diferença líquida entre os produtos, o CDB/LC "
            f"merece análise prioritária no cenário fiscal anual."
        )

        document.add_paragraph(
            "A decisão final deve separar a eficiência da aplicação isolada "
            "da eficiência fiscal anual. LCI e LCA podem continuar superiores "
            "quando a prioridade é rentabilidade líquida isenta. Produtos "
            "tributados, como CDB/LC, Tesouro Selic e Fundo DI, podem ganhar "
            "relevância quando o imposto retido no produto reduz eventual saldo "
            "adicional de tributação mínima. A aplicação prática depende da "
            "situação fiscal completa do cliente e deve ser validada por "
            "contador ou especialista tributário."
        )

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream